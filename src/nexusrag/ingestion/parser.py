"""Document parsing for PDF, DOCX, TXT, and Markdown files."""

import hashlib
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentParseError(ValueError):
    """A document cannot be parsed for a reason the user can act on."""


@dataclass
class Section:
    """A logical section within a document."""

    title: str
    content: str
    level: int = 0
    page_number: int | None = None


@dataclass
class ParsedDocument:
    """Parsed document with extracted content and metadata."""

    id: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    sections: list[Section] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return len(self.content.split())


class DocumentParser:
    """Unified parser for multiple document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported format: {extension}")

        parser_map = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt": self._parse_text,
            ".md": self._parse_markdown,
        }

        content, sections = parser_map[extension](path)
        doc_id = self._generate_id(path, content)

        return ParsedDocument(
            id=doc_id,
            content=content,
            metadata=self._extract_metadata(path),
            sections=sections,
        )

    def parse_bytes(self, data: bytes, filename: str, extension: str) -> ParsedDocument:
        """Parse document from bytes (for file uploads)."""
        import tempfile

        extension = extension.lower()
        if not extension.startswith("."):
            extension = f".{extension}"

        with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)

        try:
            doc = self.parse(tmp_path)
            # keep the user's filename, not the temp one
            doc.metadata["filename"] = filename
            doc.metadata["original_filename"] = filename
            doc.metadata["file_type"] = extension.lstrip(".")
            doc.metadata["display_name"] = filename
            return doc
        finally:
            tmp_path.unlink()

    def _parse_pdf(self, path: Path) -> tuple[str, list[Section]]:
        """Extract text and sections from PDF."""
        reader = PdfReader(path)
        if reader.is_encrypted:
            raise DocumentParseError(
                "PDF is password-protected; remove the password and upload again"
            )

        sections: list[Section] = []
        full_text_parts: list[str] = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                full_text_parts.append(text)
                page_sections = self._extract_sections_from_text(text, page_num)
                sections.extend(page_sections)

        full_text = "\n\n".join(full_text_parts)
        if not full_text.strip() and self._has_images(reader):
            raise DocumentParseError(
                "PDF contains no extractable text but has images — it is likely "
                "a scanned document; OCR is not supported"
            )

        # If no sections found, treat entire content as one section
        if not sections and full_text.strip():
            sections = [Section(title="Content", content=full_text, level=0)]

        return self._clean_text(full_text), sections

    def _parse_docx(self, path: Path) -> tuple[str, list[Section]]:
        """Extract text and sections from DOCX."""
        doc = DocxDocument(str(path))
        sections: list[Section] = []
        current_section: Section | None = None
        full_text_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text_parts.append(text)

            # Detect headings by style
            if para.style and para.style.name.startswith("Heading"):
                level = self._get_heading_level(para.style.name)
                if current_section:
                    sections.append(current_section)
                current_section = Section(title=text, content="", level=level)
            elif current_section:
                current_section.content += text + "\n"

        if current_section:
            sections.append(current_section)

        full_text = "\n\n".join(full_text_parts)

        if not sections and full_text.strip():
            sections = [Section(title="Content", content=full_text, level=0)]

        return self._clean_text(full_text), sections

    def _parse_text(self, path: Path) -> tuple[str, list[Section]]:
        content = path.read_text(encoding="utf-8")
        sections = self._extract_sections_from_text(content)
        return self._clean_text(content), sections

    def _parse_markdown(self, path: Path) -> tuple[str, list[Section]]:
        """Parse Markdown file, extracting headers as sections."""
        content = path.read_text(encoding="utf-8")
        sections: list[Section] = []
        current_section: Section | None = None

        for line in content.split("\n"):
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                if current_section:
                    sections.append(current_section)
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                current_section = Section(title=title, content="", level=level)
            elif current_section:
                current_section.content += line + "\n"

        if current_section:
            sections.append(current_section)

        if not sections and content.strip():
            sections = [Section(title="Content", content=content, level=0)]

        return self._clean_text(content), sections

    def _extract_sections_from_text(
        self, text: str, page_number: int | None = None
    ) -> list[Section]:
        """Heuristic section detection from plain text."""
        sections: list[Section] = []
        lines = text.split("\n")
        current_section: Section | None = None

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Heuristic: short uppercase lines or numbered sections are likely headers
            is_header = (
                (stripped.isupper() and len(stripped) < 100)
                or re.match(r"^\d+\.\s+[A-Z]", stripped)
                or re.match(r"^[IVX]+\.\s+", stripped)
            )

            if is_header:
                if current_section:
                    sections.append(current_section)
                current_section = Section(
                    title=stripped, content="", level=1, page_number=page_number
                )
            elif current_section:
                current_section.content += stripped + " "

        if current_section:
            sections.append(current_section)

        return sections

    @staticmethod
    def _has_images(reader: PdfReader) -> bool:
        """True when any page carries an image (best effort)."""
        try:
            return any(page.images for page in reader.pages)
        except Exception:
            return False

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from DOCX style name."""
        match = re.search(r"(\d+)", style_name)
        return int(match.group(1)) if match else 1

    def _clean_text(self, text: str) -> str:
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    def _generate_id(self, path: Path, content: str) -> str:
        """Stable id from name + length + head of content."""
        hash_input = f"{path.name}:{len(content)}:{content[:500]}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]

    def _extract_metadata(self, path: Path) -> dict[str, Any]:
        stat = path.stat()
        return {
            "filename": path.name,
            "extension": path.suffix.lower(),
            "size_bytes": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime, tz=UTC).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=UTC).isoformat(),
        }
